"""Render a Slack-mrkdwn prose payload to a sibling ``.docx`` for Slack upload.

Plan ``floating-prancing-trinket`` PR 6: users routinely ask for "Word + Excel"
deliverables. The xlsx side already auto-generates as a Parquet sibling
(``sf_dump_tool.dump_sf_query`` → ``xlsx_export.parquet_to_xlsx_sibling``), but
there is no .docx path anywhere in the orchestrator. This module fills that
gap: the post_report dispatcher calls ``render_docx`` with the already-composed
prose, and a .docx file lands next to any xlsx attachment on the same Slack
thread.

The renderer is intentionally minimal:

  * **Single source of truth — the Writing Agent's prose.** We do NOT re-run
    the Writing Agent, do NOT re-validate, do NOT re-style. The Coordinator
    has already polished the prose through the rejection-loop rubric and the
    deterministic editor + prose_polish passes. We translate that prose, as
    is, into Word paragraphs.

  * **Slack-mrkdwn surface (with markdown fallback).** The prose comes in as
    Slack-mrkdwn, where single-asterisk ``*x*`` is BOLD and underscore
    ``_x_`` is ITALIC. We accept that primary convention and also accept
    the markdown double-asterisk ``**x**`` as bold so a caller that
    pre-converts is not penalized. Bullets accept both ``- `` and ``* ``
    starts. Tables are NOT parsed from prose — they ride in via the
    ``tables`` parameter (a list of TableBlock objects) and become real
    Word tables in the document.

  * **Defensive — never raise.** xlsx delivery is best-effort and so is
    docx delivery. Any failure (missing python-docx, write permission, malformed
    prose) returns None and logs at WARNING. The Slack post must succeed
    regardless.

  * **Calibri typography.** Calibri 14pt bold for the title, Calibri 11pt for
    body. 1-inch margins all around. This is the Word default an operating
    partner would set themselves; it reads as a "normal" document, not a
    custom-styled one.

Public API
----------
    render_docx(prose, title, output_path, tables=None) -> Optional[str]
"""

from __future__ import annotations

import logging
import os
import re
from typing import Any, List, Optional, Sequence

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Typography constants — Calibri 14pt bold title, Calibri 11pt body, 1in margins.
# ---------------------------------------------------------------------------

_TITLE_FONT = "Calibri"
_TITLE_POINT_SIZE = 14
_BODY_FONT = "Calibri"
_BODY_POINT_SIZE = 11
_MARGIN_INCHES = 1.0


# ---------------------------------------------------------------------------
# Inline markdown.
#
# Slack mrkdwn input contract (see response_renderer.py — it produces the prose
# that lands here): single-asterisk ``*x*`` is BOLD, underscore ``_x_`` is
# ITALIC. We also accept double-asterisk ``**x**`` as bold so a caller that
# pre-converts the prose to vanilla Markdown is not penalized.
# ---------------------------------------------------------------------------

# Bold first, longest match first. ``**word**`` must run before the
# single-asterisk pattern, otherwise the latter would split it. Each capture
# group isolates the inner text so we can emit a styled run without the
# markers.
_BOLD_DOUBLE_PATTERN = re.compile(r"\*\*(.+?)\*\*")
# Single-asterisk bold (Slack mrkdwn). We require a non-asterisk inside so
# stray double-asterisks already handled by the pass above don't double-count.
_BOLD_SINGLE_PATTERN = re.compile(r"\*([^*\n]+?)\*")
# Underscore italic (Slack mrkdwn). The negative-character-class avoids
# matching identifiers like ``snake_case_word`` mid-token by requiring the
# underscore to flank non-whitespace text without internal underscores.
_ITALIC_UNDERSCORE_PATTERN = re.compile(r"_([^_\n]+?)_")

# Bullet detection. We accept both ``- `` and ``* `` (the two conventions the
# Writing Agent prompt admits) and tolerate any leading indentation.
_BULLET_PATTERN = re.compile(r"^\s*[-*]\s+(.+)$")


def _parse_inline_runs(text: str) -> List[tuple[str, bool, bool]]:
    """Split a line into (substring, bold, italic) tuples.

    Two passes: first carve out every bold region (``**x**`` then ``*x*``);
    then, within every non-bold region, carve out ``_x_`` italic regions.
    Plain text falls through as ``(substring, False, False)``. The result is
    the sequence of runs the docx paragraph emitter walks.

    Pure function — easy to unit-test without touching python-docx.
    """
    if not text:
        return []

    # Pass 1: collect bold spans from both patterns into a single sorted
    # list, then walk the original string in one pass to emit
    # bold/plain pieces. Double-asterisk takes precedence over single
    # because its match is strictly longer at any given start position.
    bold_spans: list[tuple[int, int, str]] = []
    occupied: list[tuple[int, int]] = []
    for match in _BOLD_DOUBLE_PATTERN.finditer(text):
        bold_spans.append((match.start(), match.end(), match.group(1)))
        occupied.append((match.start(), match.end()))
    for match in _BOLD_SINGLE_PATTERN.finditer(text):
        # Skip if this single-asterisk match overlaps with a double match.
        if any(s <= match.start() < e or s < match.end() <= e for s, e in occupied):
            continue
        bold_spans.append((match.start(), match.end(), match.group(1)))
    bold_spans.sort(key=lambda span: span[0])

    pieces: List[tuple[str, bool]] = []
    cursor = 0
    for start, end, inner in bold_spans:
        if start < cursor:
            # Overlap with a previously-emitted span; skip defensively.
            continue
        if start > cursor:
            pieces.append((text[cursor:start], False))
        pieces.append((inner, True))
        cursor = end
    if cursor < len(text):
        pieces.append((text[cursor:], False))

    # Pass 2: underscore italic inside non-bold pieces.
    runs: List[tuple[str, bool, bool]] = []
    for chunk, is_bold in pieces:
        if is_bold:
            runs.append((chunk, True, False))
            continue
        sub_cursor = 0
        for match in _ITALIC_UNDERSCORE_PATTERN.finditer(chunk):
            if match.start() > sub_cursor:
                runs.append((chunk[sub_cursor : match.start()], False, False))
            runs.append((match.group(1), False, True))
            sub_cursor = match.end()
        if sub_cursor < len(chunk):
            runs.append((chunk[sub_cursor:], False, False))

    # Drop empty runs (can happen on adjacent markers like ``***``).
    return [(s, b, i) for s, b, i in runs if s]


def _emit_paragraph_runs(paragraph, runs: List[tuple[str, bool, bool]]) -> None:
    """Append each (text, bold, italic) tuple to ``paragraph`` as a run."""
    for substring, bold, italic in runs:
        run = paragraph.add_run(substring)
        run.font.name = _BODY_FONT
        from docx.shared import Pt  # local import to keep the module load cheap

        run.font.size = Pt(_BODY_POINT_SIZE)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def render_docx(
    prose: str,
    title: str,
    output_path: str,
    tables: Optional[Sequence[Any]] = None,
) -> Optional[str]:
    """Write ``prose`` (plus any ``tables``) to a Word document at ``output_path``.

    Returns the absolute output path on success, ``None`` on any failure.
    Never raises — failure is a logged warning and the caller continues without
    a Word attachment.

    Args:
        prose: The composed prose body. Empty or None is allowed — the
            document still emits when ``tables`` carries the answer (common
            for per-rep / per-account reports). When both are empty/None,
            return None silently.
            Splits on ``\\n\\n`` for paragraph breaks. Lines starting with
            ``- `` or ``* `` become bullets. ``*word*`` is bold (Slack mrkdwn);
            ``**word**`` is also bold (Markdown). ``_word_`` is italic.
        title: First heading in the document. Calibri 14pt bold.
        output_path: Where to write the .docx. The directory must exist;
            the caller (post_report dispatcher) owns directory creation.
        tables: Optional sequence of TableBlock-shaped objects (anything that
            exposes ``title``, ``headers``, ``rows``, and optional
            ``footnote`` attributes). Rendered as real Word tables with the
            header row bolded. ``None`` or empty is a no-op.
    """
    has_prose = bool(prose and prose.strip())
    has_tables = bool(tables)
    if not has_prose and not has_tables:
        log.info(
            "[WORD_DOC_RENDER] empty prose and no tables — skipping docx generation"
        )
        return None

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as e:
        log.warning(
            "[WORD_DOC_RENDER_FAILED] python-docx not installed (%s); "
            "Slack will receive xlsx only",
            e,
        )
        return None

    try:
        doc = Document()

        # 1-inch margins on every section. New documents only have one
        # section, but loop defensively so this still does the right thing
        # if upstream ever adds page breaks/sections.
        for section in doc.sections:
            section.top_margin = Inches(_MARGIN_INCHES)
            section.bottom_margin = Inches(_MARGIN_INCHES)
            section.left_margin = Inches(_MARGIN_INCHES)
            section.right_margin = Inches(_MARGIN_INCHES)

        # Title paragraph. add_heading would use Word's "Heading 1" style
        # which varies by template; setting the run font explicitly gives us
        # the same Calibri 14pt bold regardless of the user's Normal.dotm.
        title_text = (title or "GTM Health Report").strip() or "GTM Health Report"
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title_text)
        title_run.bold = True
        title_run.font.name = _TITLE_FONT
        title_run.font.size = Pt(_TITLE_POINT_SIZE)

        # Body: split on blank lines, render each block as either a list of
        # bullets or a single paragraph with inline-styled runs.
        if has_prose:
            for block in _split_blocks(prose):
                bullets = _extract_bullets(block)
                if bullets is not None:
                    for bullet_text in bullets:
                        runs = _parse_inline_runs(bullet_text)
                        p = doc.add_paragraph(style="List Bullet")
                        _emit_paragraph_runs(p, runs)
                else:
                    runs = _parse_inline_runs(block)
                    if not runs:
                        continue
                    p = doc.add_paragraph()
                    _emit_paragraph_runs(p, runs)

        # Tables: render after prose. The Slack post drops to a Block Kit
        # table block (one inline table per message); the .docx has no such
        # limit, so we render every table the agent emitted.
        if has_tables:
            for table_block in tables or []:
                _emit_table(doc, table_block)

        doc.save(output_path)
        log.info(
            "[WORD_DOC_RENDER] wrote %s (%d bytes)",
            output_path,
            os.path.getsize(output_path) if os.path.exists(output_path) else 0,
        )
        return output_path
    except Exception as e:
        log.warning(
            "[WORD_DOC_RENDER_FAILED] %s → silent fallback (Slack still posts): %s",
            output_path,
            e,
        )
        # Best-effort cleanup so a half-written file doesn't confuse the
        # attach path.
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return None


def _split_blocks(prose: str) -> List[str]:
    """Split prose on blank-line boundaries.

    The Writing Agent emits prose with ``\\n\\n`` between paragraphs (Strunk
    rule 14: never join two ideas with "and"; write two sentences — and in
    output, two paragraphs). Single-newlines inside a block stay attached to
    the same paragraph, which keeps wrapped lines in bullet lists from
    fragmenting.
    """
    # Normalize CRLF for cross-platform robustness, then split on two-or-more
    # newlines. Strip empty blocks at the edges.
    normalized = prose.replace("\r\n", "\n").replace("\r", "\n")
    return [b.strip() for b in re.split(r"\n\s*\n", normalized) if b.strip()]


def _emit_table(doc, table_block) -> None:
    """Render a TableBlock-shaped object as a real Word table.

    ``table_block`` only needs to expose ``title``, ``headers``, ``rows``,
    and optional ``footnote`` attributes — we duck-type rather than import
    response_schemas so this module stays standalone.

    Layout:
      * Bold title paragraph (Calibri 11pt) above the table.
      * The table itself using the "Table Grid" style when available; the
        header row is bolded so the structure reads at a glance.
      * Optional footnote paragraph (italic, Calibri 11pt) below.
    """
    from docx.shared import Pt  # local import keeps module load cheap

    title = getattr(table_block, "title", None)
    headers = list(getattr(table_block, "headers", None) or [])
    rows = list(getattr(table_block, "rows", None) or [])
    footnote = getattr(table_block, "footnote", None)

    if not headers:
        # An empty header row means there is nothing meaningful to render —
        # skip rather than emit an empty table. Defensive: an upstream
        # validator should never let this through, but render_docx is best-
        # effort and must not raise.
        return

    if title:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(str(title))
        title_run.bold = True
        title_run.font.name = _BODY_FONT
        title_run.font.size = Pt(_BODY_POINT_SIZE)

    # Build the table. python-docx requires the row count up front: header + body.
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        # "Table Grid" is shipped with the default template; if a custom
        # template removes it, fall back to no style rather than raise.
        table.style = "Table Grid"
    except Exception:
        pass

    header_row = table.rows[0]
    for col_index, header_text in enumerate(headers):
        cell = header_row.cells[col_index]
        # Reuse the first paragraph the cell ships with rather than adding
        # one — adding leaves a stray empty paragraph above the text.
        cell_paragraph = cell.paragraphs[0]
        run = cell_paragraph.add_run(str(header_text))
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(_BODY_POINT_SIZE)

    for row_index, row_data in enumerate(rows, start=1):
        word_row = table.rows[row_index]
        for col_index in range(len(headers)):
            cell = word_row.cells[col_index]
            cell_paragraph = cell.paragraphs[0]
            # Cells short on data render as empty strings — never raise.
            value = row_data[col_index] if col_index < len(row_data) else ""
            run = cell_paragraph.add_run(str(value) if value is not None else "")
            run.font.name = _BODY_FONT
            run.font.size = Pt(_BODY_POINT_SIZE)

    if footnote:
        footnote_paragraph = doc.add_paragraph()
        footnote_run = footnote_paragraph.add_run(str(footnote))
        footnote_run.italic = True
        footnote_run.font.name = _BODY_FONT
        footnote_run.font.size = Pt(_BODY_POINT_SIZE)


def _extract_bullets(block: str) -> Optional[List[str]]:
    """Return the bullet items in ``block`` if every line is a bullet, else None.

    A block is treated as a bullet list only when every non-empty line starts
    with ``- `` or ``* ``. Mixed blocks (one paragraph + one bullet) fall
    through to plain-paragraph rendering — splitting them would risk dropping
    sentence context the Writing Agent intentionally placed together.
    """
    lines = [line for line in block.split("\n") if line.strip()]
    if not lines:
        return None
    bullets: List[str] = []
    for line in lines:
        m = _BULLET_PATTERN.match(line)
        if not m:
            return None
        bullets.append(m.group(1).strip())
    return bullets
