"""Integration test for the .docx auto-attach path in ``_dispatch_post_report``.

Plan ``floating-prancing-trinket`` PR 6 wires ``word_doc_renderer.render_docx``
into the post_report dispatcher so every user-facing report ships a Word
sibling alongside the xlsx. These tests verify the wiring without touching
Slack or python-docx internals beyond what the dispatcher orchestrates.

Run::

    cd orchestrator && python3 -m pytest dispatch_word_doc_attach_test.py -v
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from unittest.mock import patch


# ──────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────


def _valid_adhoc_tool_input() -> dict:
    """Minimal AdHocInvestigationResponse payload that passes Pydantic."""
    return {
        "response_type": "ad_hoc_investigation_result",
        "payload": {
            "headline": "Q3 pipeline coverage is 1.2x — below the 1.5x target",
            "key_metrics": [
                {
                    "name": "Pipeline coverage",
                    "current": "1.2x",
                    "prior": "1.6x",
                    "benchmark": "1.5x",
                    "trend": "down",
                }
            ],
            "findings": [
                {
                    "headline": "Two reps in Northeast missing 65% of the gap",
                    "value": "$4.2M short of plan",
                    "confidence": "HIGH",
                    "severity": "critical",
                }
            ],
        },
    }


def _valid_quick_answer_tool_input() -> dict:
    return {
        "response_type": "quick_answer",
        "payload": {
            "metric": "Active sales reps (quota-carrying)",
            "value": "28 total: 14 NB reps, 14 Account Managers",
            "as_of": "2026-05-14",
            "source": "Salesforce User",
        },
    }


# ──────────────────────────────────────────────────────────────────────────
# Tests
# ──────────────────────────────────────────────────────────────────────────


def test_post_report_appends_docx_to_attach_list(tmp_path: Path, monkeypatch):
    """``_dispatch_post_report`` must call render_docx and include the .docx
    in the list handed to ``_attach_files_async``."""
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)
        captured["reply_to"] = reply_to

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "1234567890.000100"
        result_text = _dispatch_post_report(
            _valid_adhoc_tool_input(),
            thread_ts="thread-x",
            session_id="sess-docx-1",
        )

    result = json.loads(result_text)
    assert result["ok"] is True

    files = captured.get("files") or []
    docx_files = [f for f in files if f.endswith(".docx")]
    assert docx_files, (
        f"expected at least one .docx attached to post_report; got: {files}"
    )

    # The docx must exist on disk and be under SESSION_OUTPUT_DIR so the
    # safe-attachment whitelist would accept it.
    docx_path = docx_files[0]
    assert os.path.exists(docx_path), f"docx file missing: {docx_path}"
    assert os.path.getsize(docx_path) > 0
    assert docx_path.startswith(str(tmp_path)), (
        f"docx must live under SESSION_OUTPUT_DIR, got: {docx_path}"
    )


def test_post_report_docx_filename_follows_sibling_pattern(tmp_path: Path, monkeypatch):
    """File name must follow ``report_<ts>_<short_uuid>.docx`` — the same
    timestamp+short-id shape used by ``sf_dump_tool`` for Parquet/xlsx siblings.
    """
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "ts"
        _dispatch_post_report(
            _valid_adhoc_tool_input(), thread_ts="t", session_id="sess-pattern"
        )

    docx_files = [f for f in captured.get("files") or [] if f.endswith(".docx")]
    assert docx_files
    name = os.path.basename(docx_files[0])
    # ``report_YYYYMMDDTHHMMSSffffff_xxxx.docx`` — same shape as sf_dump_tool's
    # ``sf_<label>_<ts>_<uuid4[:4]>.parquet``.
    assert name.startswith("report_"), name
    assert name.endswith(".docx"), name


def test_post_report_docx_title_uses_headline(tmp_path: Path, monkeypatch):
    """The Word doc's title heading must be the validated payload's headline.

    This is the readable hook the operator sees when they open the file —
    matching the Slack post's first bolded line keeps the two surfaces in
    sync without forcing the operator to scan the document body.
    """
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "ts"
        _dispatch_post_report(
            _valid_adhoc_tool_input(), thread_ts="t", session_id="sess-title"
        )

    docx_path = next(
        (f for f in captured.get("files") or [] if f.endswith(".docx")), None
    )
    assert docx_path is not None

    from docx import Document

    doc = Document(docx_path)
    headline = "Q3 pipeline coverage is 1.2x — below the 1.5x target"
    assert headline in doc.paragraphs[0].text


def test_post_report_quick_answer_uses_metric_as_title(tmp_path: Path, monkeypatch):
    """QuickAnswerResponse has no ``headline``; the title falls back to
    ``metric`` so the Word doc still has a meaningful heading."""
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "ts"
        _dispatch_post_report(
            _valid_quick_answer_tool_input(),
            thread_ts="t",
            session_id="sess-qa",
        )

    docx_path = next(
        (f for f in captured.get("files") or [] if f.endswith(".docx")), None
    )
    assert docx_path is not None

    from docx import Document

    doc = Document(docx_path)
    assert "Active sales reps" in doc.paragraphs[0].text


def test_post_report_docx_failure_does_not_block_slack_post(
    tmp_path: Path, monkeypatch
):
    """If render_docx returns None (any failure), the Slack post still goes
    through and the xlsx-only attachment path is unchanged."""
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
        patch("word_doc_renderer.render_docx", return_value=None) as mock_render,
    ):
        mock_send.return_value = "1234567890.000100"
        result_text = _dispatch_post_report(
            _valid_adhoc_tool_input(),
            thread_ts="thread-x",
            session_id="sess-docx-fail",
        )

    result = json.loads(result_text)
    # Slack post still succeeded.
    assert result["ok"] is True
    # render_docx WAS called (the wiring is in place).
    assert mock_render.called, "render_docx should be invoked for every post_report"
    # No .docx in the attachment list because render returned None.
    files = captured.get("files") or []
    docx_files = [f for f in files if f.endswith(".docx")]
    assert docx_files == [], (
        f"render_docx=None must NOT produce a docx attachment; got: {files}"
    )


def test_post_report_docx_includes_table_rows(tmp_path: Path, monkeypatch):
    """When the payload includes ``tables``, the .docx must contain every row.

    Pre-fix bug: the dispatcher only handed the rendered Slack mrkdwn string
    to render_docx; TableBlocks live in extra Block Kit blocks, so the
    Word file silently dropped per-rep / per-account row data.
    """
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    from session_runner import _dispatch_post_report

    tool_input = _valid_adhoc_tool_input()
    tool_input["payload"]["tables"] = [
        {
            "title": "Reps below quota",
            "headers": ["Rep", "Quota %", "Gap"],
            "rows": [
                ["Alice", "62%", "$120K"],
                ["Bob", "71%", "$85K"],
            ],
            "footnote": "n=2 reps, Q3 to date",
        }
    ]

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "ts"
        _dispatch_post_report(tool_input, thread_ts="t", session_id="sess-table")

    docx_path = next(
        (f for f in captured.get("files") or [] if f.endswith(".docx")), None
    )
    assert docx_path is not None, "no docx attached"

    from docx import Document

    doc = Document(docx_path)
    assert len(doc.tables) == 1, (
        f"expected the TableBlock to land as a real Word table; got "
        f"{len(doc.tables)} tables"
    )
    word_table = doc.tables[0]
    # Header + 2 rows.
    assert len(word_table.rows) == 3
    flat = {c.text for r in word_table.rows for c in r.cells}
    # Every header cell and every data cell present.
    assert {
        "Rep",
        "Quota %",
        "Gap",
        "Alice",
        "62%",
        "$120K",
        "Bob",
        "71%",
        "$85K",
    } <= flat


def test_post_report_docx_alongside_existing_attachments(tmp_path: Path, monkeypatch):
    """When the validated payload already lists attachments, the .docx is
    appended — both attachments reach ``_attach_files_async``."""
    monkeypatch.setenv("SESSION_OUTPUT_DIR", str(tmp_path))

    # Create a fake xlsx under SESSION_OUTPUT_DIR so the safe-path filter
    # accepts it.
    fake_xlsx = tmp_path / "leads_dump.xlsx"
    fake_xlsx.write_bytes(b"PK\x03\x04fake-xlsx-bytes")

    from session_runner import _dispatch_post_report

    tool_input = _valid_adhoc_tool_input()
    tool_input["payload"]["attachments"] = [str(fake_xlsx)]

    captured: dict = {}

    def _capture_attach(files, reply_to, channel=None):
        captured["files"] = list(files)

    with (
        patch("session_runner.send_notification") as mock_send,
        patch("session_runner._attach_files_async", side_effect=_capture_attach),
    ):
        mock_send.return_value = "ts"
        _dispatch_post_report(tool_input, thread_ts="t", session_id="sess-both")

    files = captured.get("files") or []
    # Both the agent-supplied xlsx AND the orchestrator-generated docx.
    assert str(fake_xlsx) in files, f"missing original xlsx in: {files}"
    docx_files = [f for f in files if f.endswith(".docx")]
    assert docx_files, f"missing auto-generated docx in: {files}"
