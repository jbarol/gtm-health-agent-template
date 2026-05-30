"""Tests for ``orchestrator.word_doc_renderer``.

Covers:
    - Happy path: render_docx writes a non-empty .docx with the title heading.
    - Defensive: None / empty / whitespace prose returns None and writes nothing.
    - Markdown surface: ``**bold**`` becomes a bold run, ``*italic*`` an
      italic run, ``- `` and ``* `` lines become a List Bullet paragraph.
    - Calibri 11pt body / Calibri 14pt bold title typography is honored.
    - Failure is never raised: a bad output_path returns None.

Run::

    cd orchestrator && python3 -m pytest word_doc_renderer_test.py -v
"""

from __future__ import annotations

from pathlib import Path


from word_doc_renderer import (
    _emit_paragraph_runs,
    _extract_bullets,
    _parse_inline_runs,
    _split_blocks,
    render_docx,
)


# ──────────────────────────────────────────────────────────────────────────
# Inline parser unit tests — pure, no python-docx dependency
# ──────────────────────────────────────────────────────────────────────────


def test_parse_inline_runs_plain_text():
    runs = _parse_inline_runs("simple plain text")
    assert runs == [("simple plain text", False, False)]


def test_parse_inline_runs_double_asterisk_bold_run():
    # ``**word**`` is markdown bold — accepted as a forgiving fallback for
    # any caller that pre-converts Slack mrkdwn to vanilla Markdown.
    runs = _parse_inline_runs("ARR is **$5.2M** this quarter")
    assert runs == [
        ("ARR is ", False, False),
        ("$5.2M", True, False),
        (" this quarter", False, False),
    ]


def test_parse_inline_runs_slack_single_asterisk_is_bold():
    # The Slack mrkdwn the response_renderer emits uses *single-asterisk*
    # for bold. This is the production input contract: every report flowing
    # through render_docx today has its headlines wrapped this way.
    runs = _parse_inline_runs("Note: *Recommended:* lower discount discipline")
    assert ("Recommended:", True, False) in runs
    # Prefix and suffix stay plain.
    assert ("Note: ", False, False) in runs
    assert (" lower discount discipline", False, False) in runs


def test_parse_inline_runs_underscore_is_italic():
    # Slack mrkdwn italic is ``_word_``. The renderer must surface this as
    # italic — not bold, not plain.
    runs = _parse_inline_runs("note _this caveat_ matters")
    assert runs == [
        ("note ", False, False),
        ("this caveat", False, True),
        (" matters", False, False),
    ]


def test_parse_inline_runs_bold_and_italic_combined():
    # Slack convention: ``*x*`` bold, ``_y_`` italic.
    runs = _parse_inline_runs("*bold* and _italic_ in one line")
    assert ("bold", True, False) in runs
    assert ("italic", False, True) in runs
    assert (" and ", False, False) in runs
    assert (" in one line", False, False) in runs


def test_parse_inline_runs_empty_string():
    assert _parse_inline_runs("") == []


def test_split_blocks_paragraph_breaks():
    prose = "First paragraph.\n\nSecond paragraph.\n\nThird."
    assert _split_blocks(prose) == [
        "First paragraph.",
        "Second paragraph.",
        "Third.",
    ]


def test_split_blocks_crlf_normalized():
    prose = "Para one.\r\n\r\nPara two."
    assert _split_blocks(prose) == ["Para one.", "Para two."]


def test_extract_bullets_dash_form():
    block = "- first item\n- second item\n- third item"
    assert _extract_bullets(block) == ["first item", "second item", "third item"]


def test_extract_bullets_asterisk_form():
    block = "* one\n* two"
    assert _extract_bullets(block) == ["one", "two"]


def test_extract_bullets_mixed_block_returns_none():
    # If a block mixes prose + bullets, treat it as prose — splitting risks
    # breaking sentence context the Writing Agent placed together.
    block = "Setup sentence.\n- bullet item\n- another"
    assert _extract_bullets(block) is None


def test_extract_bullets_empty_block_returns_none():
    assert _extract_bullets("") is None
    assert _extract_bullets("   ") is None


# ──────────────────────────────────────────────────────────────────────────
# Defensive: empty / None / whitespace prose returns None silently
# ──────────────────────────────────────────────────────────────────────────


def test_render_docx_empty_prose_returns_none(tmp_path: Path):
    out = tmp_path / "report.docx"
    result = render_docx(prose="", title="Title", output_path=str(out))
    assert result is None
    assert not out.exists()


def test_render_docx_none_prose_returns_none(tmp_path: Path):
    out = tmp_path / "report.docx"
    result = render_docx(prose=None, title="Title", output_path=str(out))  # type: ignore[arg-type]
    assert result is None
    assert not out.exists()


def test_render_docx_whitespace_only_prose_returns_none(tmp_path: Path):
    out = tmp_path / "report.docx"
    result = render_docx(prose="   \n  \n\t", title="Title", output_path=str(out))
    assert result is None
    assert not out.exists()


# ──────────────────────────────────────────────────────────────────────────
# Failure paths: never raises
# ──────────────────────────────────────────────────────────────────────────


def test_render_docx_unwritable_path_returns_none(tmp_path: Path):
    # A path under a non-existent directory will fail at doc.save time —
    # never raises, just returns None and logs.
    bad_path = tmp_path / "nonexistent_subdir" / "report.docx"
    result = render_docx(prose="some prose", title="Title", output_path=str(bad_path))
    assert result is None


# ──────────────────────────────────────────────────────────────────────────
# Happy path: writes a non-empty .docx
# ──────────────────────────────────────────────────────────────────────────


def test_render_docx_writes_non_empty_file(tmp_path: Path):
    out = tmp_path / "report.docx"
    prose = (
        "GTM health summary for last week.\n\n"
        "Pipeline coverage is at 1.2x — below the 1.5x target."
    )
    result = render_docx(prose=prose, title="GTM Weekly Briefing", output_path=str(out))
    assert result == str(out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_render_docx_emits_title_heading(tmp_path: Path):
    out = tmp_path / "report.docx"
    render_docx(
        prose="Body paragraph.",
        title="Q3 New Business Brief",
        output_path=str(out),
    )
    from docx import Document

    doc = Document(str(out))
    # First paragraph is the title; subsequent paragraphs are body.
    first = doc.paragraphs[0]
    assert "Q3 New Business Brief" in first.text
    # Title run should be bold.
    assert any(run.bold for run in first.runs if run.text), (
        f"expected bold title run, got runs: {[(r.text, r.bold) for r in first.runs]}"
    )


def test_render_docx_falls_back_to_default_title_when_missing(tmp_path: Path):
    out = tmp_path / "report.docx"
    render_docx(prose="Body.", title="", output_path=str(out))
    from docx import Document

    doc = Document(str(out))
    # Empty title falls back to a sensible default — file still readable.
    assert doc.paragraphs[0].text.strip(), (
        "expected a non-empty fallback title even when caller passed ''"
    )


# ──────────────────────────────────────────────────────────────────────────
# Markdown-style bold / italic / bullets emit the right docx runs
# ──────────────────────────────────────────────────────────────────────────


def test_render_docx_double_asterisk_becomes_bold_run(tmp_path: Path):
    # ``**x**`` (vanilla markdown bold) is honored as a fallback path.
    out = tmp_path / "report.docx"
    render_docx(
        prose="ARR climbed to **$5.2M** this quarter.",
        title="Test",
        output_path=str(out),
    )
    from docx import Document

    doc = Document(str(out))
    body_runs = []
    # Skip the title paragraph (index 0).
    for paragraph in doc.paragraphs[1:]:
        for run in paragraph.runs:
            body_runs.append((run.text, bool(run.bold), bool(run.italic)))
    # The "$5.2M" substring must be in a bold run; the surrounding text must not be.
    assert any(text == "$5.2M" and bold for text, bold, _ in body_runs), (
        f"expected '$5.2M' bold run, got: {body_runs}"
    )
    assert any("ARR climbed to" in text and not bold for text, bold, _ in body_runs), (
        f"expected non-bold prefix run, got: {body_runs}"
    )


def test_render_docx_slack_single_asterisk_becomes_bold_run(tmp_path: Path):
    # This is the production input shape — the response_renderer emits
    # Slack mrkdwn with ``*x*`` for bold. A previous version of the parser
    # treated this as italic, which broke 100% of reports' bold formatting.
    out = tmp_path / "report.docx"
    render_docx(
        prose="*Recommended:* tighten discount discipline.",
        title="Test",
        output_path=str(out),
    )
    from docx import Document

    doc = Document(str(out))
    body_runs = []
    for paragraph in doc.paragraphs[1:]:
        for run in paragraph.runs:
            body_runs.append((run.text, bool(run.bold), bool(run.italic)))
    assert any(
        text == "Recommended:" and bold and not italic
        for text, bold, italic in body_runs
    ), f"expected '*Recommended:*' as a bold (not italic) run, got: {body_runs}"


def test_render_docx_underscore_marker_becomes_italic_run(tmp_path: Path):
    # Slack mrkdwn italic uses ``_word_``.
    out = tmp_path / "report.docx"
    render_docx(
        prose="Note _this caveat_ on the data.",
        title="Test",
        output_path=str(out),
    )
    from docx import Document

    doc = Document(str(out))
    body_runs = []
    for paragraph in doc.paragraphs[1:]:
        for run in paragraph.runs:
            body_runs.append((run.text, bool(run.bold), bool(run.italic)))
    assert any(text == "this caveat" and italic for text, _, italic in body_runs), (
        f"expected 'this caveat' italic run, got: {body_runs}"
    )


def test_render_docx_dash_lines_become_bullet_list(tmp_path: Path):
    out = tmp_path / "report.docx"
    prose = (
        "Top three drivers:\n\n"
        "- Lower discount discipline\n"
        "- Slower SDR ramp\n"
        "- Sponsor churn in two accounts"
    )
    render_docx(prose=prose, title="Test", output_path=str(out))
    from docx import Document

    doc = Document(str(out))
    bullet_texts = [
        p.text for p in doc.paragraphs if p.style and "List Bullet" in p.style.name
    ]
    assert "Lower discount discipline" in bullet_texts
    assert "Slower SDR ramp" in bullet_texts
    assert "Sponsor churn in two accounts" in bullet_texts


def test_render_docx_asterisk_lines_become_bullet_list(tmp_path: Path):
    out = tmp_path / "report.docx"
    prose = "Action items:\n\n* Pull regional VP\n* Re-score health"
    render_docx(prose=prose, title="Test", output_path=str(out))
    from docx import Document

    doc = Document(str(out))
    bullet_texts = [
        p.text for p in doc.paragraphs if p.style and "List Bullet" in p.style.name
    ]
    assert "Pull regional VP" in bullet_texts
    assert "Re-score health" in bullet_texts


# ──────────────────────────────────────────────────────────────────────────
# Typography: Calibri 11pt body, Calibri 14pt bold title, 1in margins
# ──────────────────────────────────────────────────────────────────────────


def test_render_docx_typography_calibri(tmp_path: Path):
    out = tmp_path / "report.docx"
    render_docx(
        prose="Plain body paragraph here.",
        title="Title here",
        output_path=str(out),
    )
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document(str(out))

    # Title typography: Calibri 14pt bold.
    title_run = doc.paragraphs[0].runs[0]
    assert title_run.font.name == "Calibri"
    assert title_run.font.size == Pt(14)
    assert title_run.bold is True

    # Body typography: Calibri 11pt.
    body_run = doc.paragraphs[1].runs[0]
    assert body_run.font.name == "Calibri"
    assert body_run.font.size == Pt(11)

    # Section margins.
    section = doc.sections[0]
    assert section.top_margin == Inches(1.0)
    assert section.bottom_margin == Inches(1.0)
    assert section.left_margin == Inches(1.0)
    assert section.right_margin == Inches(1.0)


# ──────────────────────────────────────────────────────────────────────────
# Module-level _emit_paragraph_runs sanity check
# ──────────────────────────────────────────────────────────────────────────


def test_emit_paragraph_runs_writes_each_tuple(tmp_path: Path):
    """Ensures the helper writes one run per (text, bold, italic) tuple."""
    from docx import Document

    doc = Document()
    paragraph = doc.add_paragraph()
    _emit_paragraph_runs(
        paragraph,
        [
            ("plain ", False, False),
            ("bold", True, False),
            (" mid ", False, False),
            ("italic", False, True),
        ],
    )
    # 4 runs, in order, with the right style flags.
    runs = paragraph.runs
    assert len(runs) == 4
    assert runs[0].text == "plain " and not runs[0].bold and not runs[0].italic
    assert runs[1].text == "bold" and runs[1].bold and not runs[1].italic
    assert runs[2].text == " mid " and not runs[2].bold and not runs[2].italic
    assert runs[3].text == "italic" and not runs[3].bold and runs[3].italic


# ──────────────────────────────────────────────────────────────────────────
# Tables — TableBlock-shaped inputs become real Word tables
# ──────────────────────────────────────────────────────────────────────────


class _FakeTableBlock:
    """Duck-typed stand-in for response_schemas.TableBlock.

    Lets the renderer tests stay light — no Pydantic import, no schema
    constraints to satisfy. ``word_doc_renderer._emit_table`` only reads
    ``title``, ``headers``, ``rows``, and ``footnote`` attributes.
    """

    def __init__(self, title, headers, rows, footnote=None):
        self.title = title
        self.headers = headers
        self.rows = rows
        self.footnote = footnote


def test_render_docx_emits_table_when_provided(tmp_path: Path):
    """When ``tables=`` is non-empty, every cell appears in a real Word table.

    The pre-fix renderer dropped table data entirely because the prose
    string never carried it (the Slack renderer returns tables as separate
    Block Kit blocks). This test pins the fix.
    """
    out = tmp_path / "report.docx"
    table = _FakeTableBlock(
        title="Reps below quota",
        headers=["Rep", "Quota %", "Gap"],
        rows=[
            ["Alice", "62%", "$120K"],
            ["Bob", "71%", "$85K"],
        ],
        footnote="n=2 reps, Q3 to date",
    )
    result = render_docx(
        prose="Two reps below quota are driving the gap.",
        title="Q3 review",
        output_path=str(out),
        tables=[table],
    )
    assert result == str(out)

    from docx import Document

    doc = Document(str(out))
    # python-docx exposes added tables as Document.tables.
    assert len(doc.tables) == 1, f"expected one Word table, got {len(doc.tables)}"
    word_table = doc.tables[0]
    # Header row + 2 data rows.
    assert len(word_table.rows) == 3
    # Every header cell present.
    header_texts = [c.text for c in word_table.rows[0].cells]
    assert header_texts == ["Rep", "Quota %", "Gap"], header_texts
    # Every data cell present.
    row1 = [c.text for c in word_table.rows[1].cells]
    row2 = [c.text for c in word_table.rows[2].cells]
    assert row1 == ["Alice", "62%", "$120K"], row1
    assert row2 == ["Bob", "71%", "$85K"], row2

    # Header runs are bold.
    header_runs = word_table.rows[0].cells[0].paragraphs[0].runs
    assert any(r.bold for r in header_runs), "expected header row text to render bold"

    # Footnote appears as a paragraph below the table.
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "n=2 reps, Q3 to date" in all_text


def test_render_docx_emits_multiple_tables(tmp_path: Path):
    """The .docx is not limited to one table the way the Slack post is."""
    out = tmp_path / "report.docx"
    t1 = _FakeTableBlock("First", ["A"], [["1"]])
    t2 = _FakeTableBlock("Second", ["B"], [["2"]])
    render_docx(
        prose="See tables below.",
        title="Test",
        output_path=str(out),
        tables=[t1, t2],
    )
    from docx import Document

    doc = Document(str(out))
    assert len(doc.tables) == 2


def test_render_docx_tables_only_no_prose(tmp_path: Path):
    """A table-only payload (empty prose) still produces a Word document.

    Per-rep / per-account answers carry their answer in the table; the prose
    can be empty if the Writing Agent decided the headline + table say it all.
    """
    out = tmp_path / "report.docx"
    table = _FakeTableBlock("Reps", ["Rep"], [["Alice"]])
    result = render_docx(
        prose="",
        title="Reps",
        output_path=str(out),
        tables=[table],
    )
    assert result == str(out)
    from docx import Document

    doc = Document(str(out))
    assert len(doc.tables) == 1


def test_render_docx_empty_tables_returns_none_when_no_prose(tmp_path: Path):
    """An empty ``tables`` list AND empty prose returns None — nothing to render."""
    out = tmp_path / "report.docx"
    result = render_docx(prose="", title="Test", output_path=str(out), tables=[])
    assert result is None
    assert not out.exists()


def test_render_docx_table_with_empty_headers_skipped(tmp_path: Path):
    """Malformed TableBlock (no headers) is silently skipped, prose still renders."""
    out = tmp_path / "report.docx"
    bad_table = _FakeTableBlock("No headers", [], [])
    result = render_docx(
        prose="Body still present.",
        title="Test",
        output_path=str(out),
        tables=[bad_table],
    )
    assert result == str(out)
    from docx import Document

    doc = Document(str(out))
    # Body rendered; no table emitted because headers was empty.
    assert "Body still present." in "\n".join(p.text for p in doc.paragraphs)
    assert len(doc.tables) == 0
